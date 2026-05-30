"""Document Ingestion Parser.

Ingests PDF and DOCX files, extracts text with layout/coordinate
mappings, and determines section boundaries.

Technologies: pdfplumber, python-docx
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Literal

import pdfplumber
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.common.exceptions import DocumentIngestionError, UnsupportedDocumentTypeError
from src.ingestion.schemas import LayoutElement, ParsedDocument, ParsedPage

_LINE_TOLERANCE = 3.0
_LINE_WORD_GAP_TOLERANCE = 96.0
_COLUMN_ALIGNMENT_TOLERANCE = 48.0


@dataclass
class _Line:
    """A horizontal line of PDF words."""

    top: float
    words: list[dict[str, Any]] = field(default_factory=list)

    @property
    def x0(self) -> float:
        return min(float(word["x0"]) for word in self.words)

    @property
    def x1(self) -> float:
        return max(float(word["x1"]) for word in self.words)


@dataclass
class _Column:
    """A group of lines aligned into one reading-order column."""

    lines: list[_Line] = field(default_factory=list)

    @property
    def x0(self) -> float:
        return min(line.x0 for line in self.lines)

    @property
    def x1(self) -> float:
        return max(line.x1 for line in self.lines)


class DocumentIngestor:
    """Parse supported source documents into a normalized representation."""

    def parse(self, source_path: str | Path) -> ParsedDocument:
        """Dispatch parsing based on the source file extension."""

        path = Path(source_path)
        if not path.is_file():
            raise DocumentIngestionError(f"document does not exist: {path}")

        extension = path.suffix.casefold()
        if extension == ".pdf":
            return self._parse_pdf(path)
        if extension == ".docx":
            return self._parse_docx(path)
        raise UnsupportedDocumentTypeError(f"unsupported document type: {path.suffix or '<none>'}")

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Extract PDF pages with coordinate-backed word provenance."""

        try:
            with pdfplumber.open(path) as pdf:
                pages = [self._parse_pdf_page(page, page_index) for page_index, page in enumerate(pdf.pages)]
        except DocumentIngestionError:
            raise
        except Exception as exc:
            raise DocumentIngestionError(f"failed to parse PDF document: {path.name}") from exc

        return ParsedDocument(
            document_name=path.name,
            source_path=str(path),
            document_type="pdf",
            raw_pages=pages,
        )

    def _parse_pdf_page(self, page: Any, page_index: int) -> ParsedPage:
        words = page.extract_words(use_text_flow=False, keep_blank_chars=False) or []
        ordered_words = _order_pdf_words(words)
        layout_info = [
            LayoutElement(
                kind="word",
                text=str(word["text"]),
                x0=float(word["x0"]),
                top=float(word["top"]),
                x1=float(word["x1"]),
                bottom=float(word["bottom"]),
            )
            for word in ordered_words
        ]
        return ParsedPage(
            page_index=page_index,
            text=_render_pdf_text(ordered_words),
            layout_info=layout_info,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Extract ordered DOCX paragraphs and tables into logical pages."""

        try:
            document = Document(str(path))
            raw_pages = _extract_docx_pages(document)
        except Exception as exc:
            raise DocumentIngestionError(f"failed to parse DOCX document: {path.name}") from exc

        return ParsedDocument(
            document_name=path.name,
            source_path=str(path),
            document_type="docx",
            raw_pages=raw_pages,
        )


def _order_pdf_words(words: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Order PDF words by column, line, and horizontal position."""

    lines = _group_words_into_lines(words)
    columns: list[_Column] = []

    for line in sorted(lines, key=lambda item: (item.x0, item.top)):
        column = next((item for item in columns if _belongs_to_column(line, item)), None)
        if column is None:
            columns.append(_Column(lines=[line]))
        else:
            column.lines.append(line)

    ordered_words: list[dict[str, Any]] = []
    for column in sorted(columns, key=lambda item: item.x0):
        for line in sorted(column.lines, key=lambda item: (item.top, item.x0)):
            ordered_words.extend(sorted(line.words, key=lambda word: float(word["x0"])))
    return ordered_words


def _group_words_into_lines(words: list[dict[str, Any]]) -> list[_Line]:
    lines: list[_Line] = []
    for word in sorted(words, key=lambda item: (float(item["top"]), float(item["x0"]))):
        top = float(word["top"])
        x0 = float(word["x0"])
        line = next(
            (
                item
                for item in lines
                if abs(item.top - top) <= _LINE_TOLERANCE and x0 - item.x1 <= _LINE_WORD_GAP_TOLERANCE
            ),
            None,
        )
        if line is None:
            lines.append(_Line(top=top, words=[word]))
        else:
            line.words.append(word)
    return lines


def _belongs_to_column(line: _Line, column: _Column) -> bool:
    horizontal_overlap = min(line.x1, column.x1) - max(line.x0, column.x0)
    return horizontal_overlap > 0 or abs(line.x0 - column.x0) <= _COLUMN_ALIGNMENT_TOLERANCE


def _render_pdf_text(words: list[dict[str, Any]]) -> str:
    """Render ordered PDF words while preserving line changes."""

    rendered_lines: list[str] = []
    current_words: list[str] = []
    current_top: float | None = None
    previous_x1: float | None = None

    for word in words:
        top = float(word["top"])
        x0 = float(word["x0"])
        starts_new_line = (
            current_top is not None
            and previous_x1 is not None
            and (
                abs(current_top - top) > _LINE_TOLERANCE
                or x0 - previous_x1 > _LINE_WORD_GAP_TOLERANCE
                or x0 < previous_x1
            )
        )
        if starts_new_line:
            rendered_lines.append(" ".join(current_words))
            current_words = []
        current_words.append(str(word["text"]))
        current_top = top
        previous_x1 = float(word["x1"])

    if current_words:
        rendered_lines.append(" ".join(current_words))
    return "\n".join(rendered_lines)


def _extract_docx_pages(document: DocxDocument) -> list[ParsedPage]:
    page_elements: list[list[LayoutElement]] = [[]]
    source_index = 0

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            kind: Literal["heading", "paragraph"] = (
                "heading" if block.style and block.style.name.casefold().startswith("heading") else "paragraph"
            )
            fragments = _split_paragraph_on_page_breaks(block)
            for fragment_index, fragment in enumerate(fragments):
                text = fragment.strip()
                if text:
                    page_elements[-1].append(LayoutElement(kind=kind, text=text, source_index=source_index))
                    source_index += 1
                if fragment_index < len(fragments) - 1:
                    page_elements.append([])
        else:
            text = _render_docx_table(block)
            if text:
                page_elements[-1].append(LayoutElement(kind="table", text=text, source_index=source_index))
                source_index += 1

    return [
        ParsedPage(
            page_index=page_index,
            text="\n".join(element.text for element in elements),
            layout_info=elements,
        )
        for page_index, elements in enumerate(page_elements)
    ]


def _iter_docx_blocks(document: DocxDocument) -> list[Paragraph | Table]:
    blocks: list[Paragraph | Table] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, document))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, document))
    return blocks


def _split_paragraph_on_page_breaks(paragraph: Paragraph) -> list[str]:
    fragments = [""]
    for node in paragraph._p.iter():  # noqa: SLF001 - python-docx exposes no ordered break-aware text API.
        if node.tag == qn("w:t"):
            fragments[-1] += node.text or ""
        elif node.tag == qn("w:tab"):
            fragments[-1] += "\t"
        elif node.tag == qn("w:br"):
            if node.get(qn("w:type")) == "page":
                fragments.append("")
            else:
                fragments[-1] += "\n"
        elif node.tag == qn("w:lastRenderedPageBreak"):
            fragments.append("")
    return fragments


def _render_docx_table(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(row for row in rows if row.strip(" |"))
