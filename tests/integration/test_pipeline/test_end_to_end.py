"""Integration tests for synchronous Phase 1 pre-flight composition."""

from pathlib import Path
from typing import Any, cast
from collections.abc import Iterable

import networkx as nx
import pytest
from docx import Document

from src.agents.planner_agent import PlannerAgent
from src.agents.schemas import ActiveSpecialistManifest, AgentName
from src.common.exceptions import DocumentIngestionError, SearchEngineError
from src.common.models import ClauseType
from src.graphrag.entity_resolver import EntityResolver
from src.graphrag.graph_constructor import GraphConstructor
from src.ingestion.chunker import SectionAwareChunker
from src.ingestion.classifier import ClauseClassifier
from src.ingestion.ingest import DocumentIngestor
from src.ingestion.schemas import DocumentChunk, ParsedDocument, ParsedPage
from src.preflight import PreflightPipeline, PreflightResult
from src.search.indexer import ElasticsearchIndexer
from src.search.schemas import IndexingOutcome


class _Ingestor:
    def __init__(self, calls: list[str]) -> None:
        self.calls = calls

    def parse(self, path: str | Path) -> ParsedDocument:
        self.calls.append(f"parse:{path}")
        return ParsedDocument(
            document_name=str(path),
            source_path=str(path),
            document_type="pdf",
            raw_pages=[ParsedPage(page_index=0, text="source")],
        )


class _Chunker:
    def __init__(self, calls: list[str]) -> None:
        self.calls = calls

    def chunk(self, document: ParsedDocument) -> list[DocumentChunk]:
        self.calls.append(f"chunk:{document.document_name}")
        return [_chunk()]


class _Classifier:
    def __init__(self, calls: list[str]) -> None:
        self.calls = calls

    def classify_chunk(self, chunk: DocumentChunk) -> DocumentChunk:
        self.calls.append(f"classify:{chunk.chunk_id}")
        return chunk.model_copy(update={"clause_type": ClauseType.IP_ASSIGNMENT})


class _GraphConstructor:
    def __init__(self, calls: list[str]) -> None:
        self.calls = calls

    def build_graph(self, chunks: list[DocumentChunk]) -> nx.DiGraph:
        self.calls.append(f"graph:{len(chunks)}")
        return nx.DiGraph()


class _Indexer:
    def __init__(self, calls: list[str], *, errors: list[dict[str, object]] | None = None) -> None:
        self.calls = calls
        self.errors = errors or []

    def ensure_index(self) -> bool:
        self.calls.append("ensure_index")
        return True

    def index_chunks(self, chunks: Iterable[DocumentChunk]) -> IndexingOutcome:
        chunks_list = list(chunks)
        self.calls.append(f"index:{len(chunks_list)}")
        return IndexingOutcome(attempted=len(chunks_list), indexed=len(chunks_list) - len(self.errors), errors=self.errors)


class _Planner:
    def __init__(self, calls: list[str]) -> None:
        self.calls = calls

    def plan(self, *, document_names: list[str], chunks: list[DocumentChunk]) -> ActiveSpecialistManifest:
        self.calls.append(f"plan:{len(document_names)}:{len(chunks)}")
        return ActiveSpecialistManifest(active_agents=[AgentName.IP])


class _Indices:
    def __init__(self) -> None:
        self.created: list[dict[str, object]] = []

    def exists(self, *, index: str) -> bool:
        return False

    def create(self, **kwargs: object) -> None:
        self.created.append(kwargs)


class _ElasticsearchClient:
    def __init__(self) -> None:
        self.indices = _Indices()


class _Entity:
    def __init__(self, text: str, label: str) -> None:
        self.text = text
        self.label_ = label


class _Document:
    def __init__(self, entities: list[_Entity]) -> None:
        self.ents = entities


class _NLP:
    def __call__(self, _text: str) -> _Document:
        return _Document(
            [
                _Entity("Acme Corp.", "ORG"),
                _Entity("Acme Corporation", "ORG"),
            ]
        )


class _Ollama:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def disambiguate(self, **kwargs: object) -> dict[str, object]:
        self.calls.append(kwargs)
        return {
            "merge": True,
            "canonical_name": "Acme Corp.",
            "confidence": 0.92,
            "reason": "Known legal alias",
        }


def test_preflight_pipeline_runs_in_required_order_and_returns_typed_output() -> None:
    calls: list[str] = []

    result = _pipeline(calls).run(["agreement.pdf"])

    assert isinstance(result, PreflightResult)
    assert calls == [
        "parse:agreement.pdf",
        "chunk:agreement.pdf",
        "classify:agreement:0",
        "graph:1",
        "ensure_index",
        "index:1",
        "plan:1:1",
    ]
    assert result.chunks[0].clause_type is ClauseType.IP_ASSIGNMENT
    assert result.specialist_manifest.active_agents == [AgentName.IP]


def test_preflight_pipeline_rejects_empty_source_manifest() -> None:
    with pytest.raises(DocumentIngestionError, match="at least one source"):
        _pipeline([]).run([])


def test_preflight_pipeline_stops_before_planner_when_indexing_is_incomplete() -> None:
    calls: list[str] = []
    pipeline = _pipeline(calls, index_errors=[{"index": {"error": "offline"}}])

    with pytest.raises(SearchEngineError, match="failed to index"):
        pipeline.run(["agreement.pdf"])

    assert calls[-1] == "index:1"
    assert not any(call.startswith("plan:") for call in calls)


def test_preflight_pipeline_processes_temporary_document_end_to_end(tmp_path: Path) -> None:
    source_path = tmp_path / "privacy-review.docx"
    source = Document()
    source.add_heading("Section 1. Privacy", level=1)
    source.add_paragraph("Acme Corp. and Acme Corporation process personal data under GDPR.")
    source.save(str(source_path))
    ollama = _Ollama()
    elasticsearch = _ElasticsearchClient()
    indexed_actions: list[dict[str, Any]] = []

    def bulk_writer(
        _client: object,
        actions: list[dict[str, Any]],
        **_kwargs: Any,
    ) -> tuple[int, list[dict[str, Any]]]:
        indexed_actions.extend(actions)
        return len(actions), []

    pipeline = PreflightPipeline(
        ingestor=DocumentIngestor(),
        chunker=SectionAwareChunker(),
        classifier=ClauseClassifier(),
        graph_constructor=GraphConstructor(
            nlp=_NLP(),
            resolver=EntityResolver(llm_client=ollama),
        ),
        indexer=ElasticsearchIndexer(
            elasticsearch,
            index_name="verdictos_documents",
            bulk_writer=bulk_writer,
        ),
        planner=PlannerAgent(),
    )

    result = pipeline.run([source_path])

    assert len(result.documents) == 1
    assert len(result.chunks) == 1
    assert result.chunks[0].clause_type is ClauseType.DATA_PROTECTION
    assert result.indexing.indexed == 1
    assert indexed_actions[0]["_source"]["clause_type"] == "data_protection"
    assert AgentName.PRIVACY in result.specialist_manifest.active_agents
    entity_nodes = [
        attributes
        for _, attributes in result.graph.nodes(data=True)
        if attributes["node_type"] == "entity"
    ]
    assert entity_nodes[0]["aliases"] == ["Acme Corp.", "Acme Corporation"]
    assert len(ollama.calls) == 1


def _pipeline(
    calls: list[str],
    *,
    index_errors: list[dict[str, object]] | None = None,
) -> PreflightPipeline:
    return PreflightPipeline(
        ingestor=cast(Any, _Ingestor(calls)),
        chunker=cast(Any, _Chunker(calls)),
        classifier=cast(Any, _Classifier(calls)),
        graph_constructor=cast(Any, _GraphConstructor(calls)),
        indexer=cast(Any, _Indexer(calls, errors=index_errors)),
        planner=cast(Any, _Planner(calls)),
    )


def _chunk() -> DocumentChunk:
    return DocumentChunk(
        chunk_id="agreement:0",
        document_name="agreement.pdf",
        chunk_index=0,
        text="source",
        section_id="Section 1",
        absolute_page=0,
    )


def test_preflight_pipeline_deduplicates_duplicate_files(tmp_path: Path) -> None:
    f1 = tmp_path / "f1.pdf"
    f1.write_bytes(b"pdf content 1")
    # f2 has the exact same content as f1
    f2 = tmp_path / "f2.pdf"
    f2.write_bytes(b"pdf content 1")
    # f3 has different content
    f3 = tmp_path / "f3.pdf"
    f3.write_bytes(b"pdf content 2")

    calls: list[str] = []
    pipeline = _pipeline(calls)
    pipeline.run([f1, f2, f3])

    # Should only parse unique files (f1 and f3, f2 is skipped as duplicate)
    parsed_files = [call for call in calls if call.startswith("parse:")]
    assert len(parsed_files) == 2
    assert f"parse:{f1}" in parsed_files
    assert f"parse:{f3}" in parsed_files
    assert f"parse:{f2}" not in parsed_files

