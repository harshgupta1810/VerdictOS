"""End-to-End Pipeline Integration Test (Phases 1-6)."""

import pytest
from unittest.mock import patch, AsyncMock, MagicMock
from pathlib import Path

from src.workers.tasks import run_deal_pipeline
from src.db.models import Deal
from src.agents.schemas import ActiveSpecialistManifest, AgentName
from src.llm.schemas import LLMResponse, LLMRequest

from docx import Document

@pytest.mark.asyncio
async def test_e2e_pipeline_full_run(tmp_path: Path):
    # Setup test file
    doc_path = tmp_path / "agreement.docx"
    doc = Document()
    doc.add_heading("Section 1", level=1)
    doc.add_paragraph("Acme Corp agrees to transfer intellectual property to Buyer Inc. No pending litigation exists.")
    doc.save(str(doc_path))
    
    # We will patch external clients (DB, ES, Ollama) but allow the core pipeline to run.
    with patch("src.workers.tasks.Elasticsearch") as mock_es, \
         patch("src.workers.tasks.OllamaClient") as mock_ollama_class, \
         patch("src.workers.tasks.AsyncSessionLocal") as mock_db_session:

        # Mock DB
        mock_session = AsyncMock()
        mock_db_session.return_value.__aenter__.return_value = mock_session
        
        # Setup Deal mock
        mock_deal = MagicMock(spec=Deal)
        mock_deal.deal_id = "deal123"
        mock_deal.status = "new"
        
        mock_session.get.return_value = mock_deal
        
        # Mock execute for DB to avoid coroutine 'all' error
        mock_db_result = MagicMock()
        mock_db_result.scalars.return_value.all.return_value = []
        mock_db_result.scalar_one_or_none.return_value = None
        mock_session.execute.return_value = mock_db_result
        
        # Mock Elasticsearch client instance
        es_instance = MagicMock()
        mock_es.return_value = es_instance
        es_instance.indices.exists.return_value = False
        es_instance.helpers = MagicMock()
        es_instance.helpers.bulk.return_value = (1, [])
        es_instance.search.return_value = {
            "hits": {
                "hits": [
                    {
                        "_id": "chunk_1",
                        "_score": 1.0,
                        "_source": {
                            "chunk_id": "chunk_1",
                            "document_name": "agreement.txt",
                            "text": "Acme Corp agrees to transfer intellectual property...",
                            "section_id": "Sec 1",
                            "absolute_page": 1,
                            "clause_type": "ip_assignment"
                        }
                    }
                ]
            }
        }
        
        # Mock Ollama client
        ollama_instance = AsyncMock()
        mock_ollama_class.return_value = ollama_instance
        ollama_instance.default_model = "llama3"
        
        # Phase 1: Planner
        manifest = ActiveSpecialistManifest(active_agents=[AgentName.IP])
        
        # Phase 2: Analysis Result
        from src.agents.base_agent import RelevanceCheckResult, _AnalysisOutput, _RawFinding
        ip_relevance = RelevanceCheckResult(relevant=True, reason="IP")
        ip_analysis = _AnalysisOutput(findings=[
            _RawFinding(claim="IP rights transferred.", citation="Acme Corp agrees to transfer intellectual property to Buyer Inc.", confidence=0.9)
        ])
        
        # Phase 3-5: Debate
        from src.debate.schemas import DebateArgument, DebatePersona, DebateStance, DimensionState
        from src.agents.schemas import FindingDimension
        arg = DebateArgument(
            id="arg_1",
            finding_id="finding_1",
            persona=DebatePersona.CRITIC,
            round=1,
            dimension=FindingDimension.RISK_EXPOSURE,
            stance=DebateStance.OPPOSE,
            argument="This is bad.",
            steelman="I understand your point",
            citations=["chunk_1"],
            dropout_flag=False
        )
        
        # Phase 6: Judge Synthesis
        from src.agents.schemas import JudgeSynthesisResult, JudgeVerdictStatus
        judge_synth = JudgeSynthesisResult(
            finding_id="finding_1",
            status=JudgeVerdictStatus.CONFIRMED,
            judge_override_flag=False,
            synthesis_rationale="Consensus reached.",
            calibrated_confidence=0.9
        )
        
        # Return different schemas based on requested Pydantic model
        async def mock_generate(prompt: LLMRequest, schema, **kwargs):
            schema_name = schema.__name__
            if schema_name == "ActiveSpecialistManifest":
                return manifest
            elif schema_name == "RelevanceCheckResult":
                return ip_relevance
            elif schema_name == "_AnalysisOutput":
                return ip_analysis
            elif schema_name == "DebateArgument":
                return arg
            elif schema_name == "JudgeSynthesisResult":
                return judge_synth
            elif schema_name == "dict": # GraphRAG entity resolution
                return {
                    "merge": True,
                    "canonical_name": "Acme Corp.",
                    "confidence": 0.92,
                    "reason": "Alias"
                }
            return MagicMock(spec=schema)
            
        ollama_instance.generate_with_schema.side_effect = mock_generate
        ollama_instance.generate.return_value = LLMResponse(model="llama3", raw_text='{"merge": true, "canonical_name": "Acme Corp.", "confidence": 0.92, "reason": "Alias"}')

        # To avoid actual graph NLP entity resolution errors, patch GraphConstructor
        with patch("src.workers.tasks.GraphConstructor.build_graph") as mock_build_graph, \
             patch("src.workers.tasks.emit_pipeline_event") as mock_emit:
            
            import networkx as nx
            mock_build_graph.return_value = nx.DiGraph()
            
            # Execute the pipeline
            await run_deal_pipeline("deal123", [str(doc_path)])
            
            # Assertions to ensure all phases completed and status progressed to judging
            # The last status transition in run_deal_pipeline is 'judging' then 'complete'
            # Note: assembler.generate_verdict updates DB and we patched emit_pipeline_event to see 'complete'
            mock_emit.assert_any_call("deal123", "phase_transition", {"phase": "complete", "progress": 100})
            
            # Verify status progression on Deal object
            assert mock_deal.status in ["judging", "complete"]
