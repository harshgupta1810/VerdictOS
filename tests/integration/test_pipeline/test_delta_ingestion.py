"""Integration tests for Delta Ingestion."""

import pytest
from pathlib import Path
from unittest.mock import MagicMock

from src.preflight import PreflightPipeline
from src.ingestion.ingest import DocumentIngestor
from src.ingestion.chunker import SectionAwareChunker
from src.ingestion.classifier import ClauseClassifier
from src.graphrag.graph_constructor import GraphConstructor
from src.agents.planner_agent import PlannerAgent
from src.search.indexer import ElasticsearchIndexer
from docx import Document

def test_delta_ingestion_creates_separate_graphrag_update(tmp_path: Path):
    # Setup first document
    doc1_path = tmp_path / "base_agreement.docx"
    doc1 = Document()
    doc1.add_heading("Section 1", level=1)
    doc1.add_paragraph("Acme Corp transfers IP to Buyer Inc.")
    doc1.save(str(doc1_path))
    
    # Setup delta document
    doc2_path = tmp_path / "delta_addendum.docx"
    doc2 = Document()
    doc2.add_heading("Section 1", level=1)
    doc2.add_paragraph("Beta Corp assumes liability for Acme Corp IP.")
    doc2.save(str(doc2_path))
    
    # Mock indexer
    from src.search.schemas import IndexingOutcome
    mock_indexer = MagicMock(spec=ElasticsearchIndexer)
    mock_indexer.index_chunks.return_value = IndexingOutcome(attempted=1, indexed=1, errors=[])
    
    # Mock Planner
    mock_planner = MagicMock(spec=PlannerAgent)
    
    # Mock GraphConstructor
    import networkx as nx
    mock_graph_constructor = MagicMock(spec=GraphConstructor)
    graph1 = nx.DiGraph()
    graph1.add_node("Acme")
    graph2 = nx.DiGraph()
    graph2.add_node("Beta")
    mock_graph_constructor.build_graph.side_effect = [graph1, graph2]
    
    pipeline = PreflightPipeline(
        ingestor=DocumentIngestor(),
        chunker=SectionAwareChunker(),
        classifier=ClauseClassifier(),
        graph_constructor=mock_graph_constructor,
        indexer=mock_indexer,
        planner=mock_planner
    )
    
    # Run pipeline on base document
    result1 = pipeline.run([str(doc1_path)])
    
    # Run pipeline on delta document
    result2 = pipeline.run([str(doc2_path)])
    
    # Verify separate GraphRAG updates
    assert mock_graph_constructor.build_graph.call_count == 2
    assert result1.graph is not result2.graph
    assert "Acme" in result1.graph.nodes
    assert "Beta" in result2.graph.nodes
