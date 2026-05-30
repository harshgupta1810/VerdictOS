"""Tests for document ingestion parser and schemas."""

from pathlib import Path

import pytest
from docx import Document
from pydantic import ValidationError

from src.common.exceptions import DocumentIngestionError, UnsupportedDocumentTypeError
from src.common.models import ClauseType
from src.ingestion.ingest import DocumentIngestor
from src.ingestion.schemas import ClassificationResult, DocumentChunk, LayoutElement, ParsedDocument, ParsedPage


def test_ingestion_schemas_accept_valid_document_chunk_and_classification() -> None:
    layout = LayoutElement(kind="word", text="Agreement", x0=1, top=2, x1=3, bottom=4)
    page = ParsedPage(page_index=0, text="Agreement", layout_info=[layout])
    document = ParsedDocument(
        document_name="agreement.pdf",
        source_path="agreement.pdf",
        document_type="pdf",
        raw_pages=[page],
    )
    chunk = DocumentChunk(
        chunk_id="agreement:0",
        document_name=document.document_name,
        chunk_index=0,
        text=document.raw_pages[0].text,
        section_id="document_root",
        absolute_page=0,
    )
    classification = ClassificationResult(
        clause_type=ClauseType.GENERAL,
        confidence=0.0,
    )

    assert chunk.clause_type is ClauseType.GENERAL
    assert classification.matched_terms == []


@pytest.mark.parametrize(
    ("model", "payload"),
    [
        (LayoutElement, {"kind": "word", "text": "bad", "x0": 1}),
        (ParsedPage, {"page_index": -1, "text": ""}),
        (
            ParsedDocument,
            {"document_name": "bad.txt", "source_path": "bad.txt", "document_type": "txt"},
        ),
        (
            DocumentChunk,
            {
                "chunk_id": "bad:0",
                "document_name": "bad.pdf",
                "chunk_index": 0,
                "text": "text",
                "section_id": "document_root",
                "absolute_page": -1,
            },
        ),
        (ClassificationResult, {"clause_type": "general", "confidence": 1.1}),
    ],
)
def test_ingestion_schemas_reject_invalid_payloads(model: type[object], payload: dict[str, object]) -> None:
    with pytest.raises(ValidationError):
        model(**payload)


def test_pdf_parser_extracts_pages_coordinates_and_column_order(tmp_path: Path) -> None:
    pdf_path = tmp_path / "columns.pdf"
    _write_pdf(
        pdf_path,
        [
            [
                (50, 750, "Left One"),
                (300, 750, "Right One"),
                (50, 730, "Left Two"),
                (300, 730, "Right Two"),
            ],
            [(50, 750, "Second Page")],
        ],
    )

    document = DocumentIngestor().parse(pdf_path)

    assert document.document_type == "pdf"
    assert [page.page_index for page in document.raw_pages] == [0, 1]
    assert document.raw_pages[0].text == "Left One\nLeft Two\nRight One\nRight Two"
    assert [word.text for word in document.raw_pages[0].layout_info] == [
        "Left",
        "One",
        "Left",
        "Two",
        "Right",
        "One",
        "Right",
        "Two",
    ]
    assert all(word.x0 is not None and word.bottom is not None for word in document.raw_pages[0].layout_info)


def test_pdf_parser_wraps_malformed_pdf_errors(tmp_path: Path) -> None:
    pdf_path = tmp_path / "bad.pdf"
    pdf_path.write_bytes(b"not a pdf")

    with pytest.raises(DocumentIngestionError, match="failed to parse PDF"):
        DocumentIngestor().parse(pdf_path)


def test_parser_rejects_missing_and_unsupported_documents(tmp_path: Path) -> None:
    with pytest.raises(DocumentIngestionError, match="does not exist"):
        DocumentIngestor().parse(tmp_path / "missing.pdf")

    text_path = tmp_path / "notes.txt"
    text_path.write_text("unsupported", encoding="utf-8")
    with pytest.raises(UnsupportedDocumentTypeError, match="unsupported document type"):
        DocumentIngestor().parse(text_path)


def test_docx_parser_extracts_headings_tables_and_logical_pages(tmp_path: Path) -> None:
    docx_path = tmp_path / "agreement.docx"
    document = Document()
    document.add_heading("Section 1. Assets", level=1)
    document.add_paragraph("The equipment is listed below.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Asset"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Server"
    table.cell(1, 1).text = "Owned"
    document.add_page_break()
    document.add_paragraph("Section 2. Privacy")
    document.save(str(docx_path))

    parsed = DocumentIngestor().parse(docx_path)

    assert parsed.document_type == "docx"
    assert [page.page_index for page in parsed.raw_pages] == [0, 1]
    assert parsed.raw_pages[0].text == (
        "Section 1. Assets\n"
        "The equipment is listed below.\n"
        "Asset | Status\n"
        "Server | Owned"
    )
    assert [element.kind for element in parsed.raw_pages[0].layout_info] == ["heading", "paragraph", "table"]
    assert parsed.raw_pages[1].text == "Section 2. Privacy"


def test_docx_parser_uses_page_zero_without_explicit_breaks(tmp_path: Path) -> None:
    docx_path = tmp_path / "single-page.docx"
    document = Document()
    document.add_paragraph("One logical page.")
    document.save(str(docx_path))

    parsed = DocumentIngestor().parse(docx_path)

    assert len(parsed.raw_pages) == 1
    assert parsed.raw_pages[0].page_index == 0


def test_docx_parser_wraps_malformed_document_errors(tmp_path: Path) -> None:
    docx_path = tmp_path / "bad.docx"
    docx_path.write_bytes(b"not a docx")

    with pytest.raises(DocumentIngestionError, match="failed to parse DOCX"):
        DocumentIngestor().parse(docx_path)


def _write_pdf(path: Path, pages: list[list[tuple[int, int, str]]]) -> None:
    """Write a small text-only PDF without adding a test dependency."""

    objects: list[bytes] = []
    page_object_numbers: list[int] = []
    content_object_numbers: list[int] = []
    font_object_number = 3 + (2 * len(pages))

    for page_index in range(len(pages)):
        page_object_numbers.append(3 + (2 * page_index))
        content_object_numbers.append(4 + (2 * page_index))

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{number} 0 R" for number in page_object_numbers)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())

    for page_object, content_object, page_lines in zip(page_object_numbers, content_object_numbers, pages, strict=True):
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_object_number} 0 R >> >> "
                f"/Contents {content_object} 0 R >>"
            ).encode()
        )
        commands = " ".join(
            f"BT /F1 12 Tf 1 0 0 1 {x} {y} Tm ({text}) Tj ET" for x, y, text in page_lines
        ).encode()
        objects.append(f"<< /Length {len(commands)} >>\nstream\n".encode() + commands + b"\nendstream")

    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_number, content in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_number} 0 obj\n".encode())
        output.extend(content)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode())
    path.write_bytes(output)


def test_compute_file_fingerprint(tmp_path: Path) -> None:
    from src.ingestion.ingest import compute_file_fingerprint
    f1 = tmp_path / "f1.txt"
    f1.write_text("hello", encoding="utf-8")
    f2 = tmp_path / "f2.txt"
    f2.write_text("hello", encoding="utf-8")
    f3 = tmp_path / "f3.txt"
    f3.write_text("world", encoding="utf-8")

    assert compute_file_fingerprint(f1) == compute_file_fingerprint(f2)
    assert compute_file_fingerprint(f1) != compute_file_fingerprint(f3)

